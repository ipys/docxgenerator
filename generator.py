"""
generator.py — Professional Academic Term Paper Builder
========================================================
Handles both:
  • Claude AI content generation
  • python-docx document assembly (cover page, body, references)
"""

import json
import re
import requests
from io import BytesIO

GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
BODY_FONT    = "Times New Roman"
HEAD_FONT    = "Times New Roman"
RED          = RGBColor(0xC0, 0x00, 0x00)
BLACK        = RGBColor(0x00, 0x00, 0x00)
DARK_BLUE    = RGBColor(0x1F, 0x39, 0x64)

# Page geometry in Cm (A4)
PAGE_W = 21.0
PAGE_H = 29.7
MARGIN_TOP    = 2.54
MARGIN_BOTTOM = 2.54
MARGIN_LEFT   = 3.00
MARGIN_RIGHT  = 2.54


# ═══════════════════════════════════════════
#  ANTHROPIC  ── Content generation
# ═══════════════════════════════════════════
def generate_paper_content(
    title: str,
    api_key: str,
    university: str = "Tikrit University",
    college: str = "College of Petroleum Process Engineering",
    department: str = "Oil and Gas Refining Department",
    grade: str = "3rd Grade",
) -> dict:
    """Call Gemini API → return structured paper content as a dict."""

    prompt = f"""You are a senior academic writer at a petroleum/chemical engineering faculty.
Write a complete, technically rigorous term paper on: "{title}"

This paper is for a student at:
  University : {university}
  College    : {college}
  Department : {department}
  Grade      : {grade}

Return ONLY valid JSON — no markdown fences, no extra text — with exactly this structure:
{{
  "introduction": "...",
  "sections": [
    {{
      "heading": "...",
      "body": "...",
      "figure_caption": "Fig 1. ..."
    }},
    {{
      "heading": "...",
      "body": "...",
      "figure_caption": "Fig 2. ..."
    }},
    {{
      "heading": "...",
      "body": "...",
      "figure_caption": null
    }},
    {{
      "heading": "...",
      "body": "...",
      "figure_caption": null
    }}
  ],
  "conclusion": "...",
  "references": [
    "APA reference 1",
    "APA reference 2",
    "APA reference 3",
    "APA reference 4",
    "APA reference 5"
  ]
}}

Rules:
- introduction : 160-200 words, no heading text inside it
- sections     : exactly 4, each body 160-210 words
  Section 1 heading must relate to the position/role of {title} in industrial systems; give it figure_caption
  Section 2 heading must relate to principle of operation; give it figure_caption
  Section 3 heading must relate to key parameters / control considerations; figure_caption = null
  Section 4 heading must relate to industrial importance / applications; figure_caption = null
- conclusion   : 110-150 words, no heading text inside it
- references   : exactly 5, proper APA 7th edition format.
  Include at least 2 with real DOI/URL links.
  Mix: at least 1 book, 1 journal article, 1 website.
- Plain text only — no bold (**), no asterisks, no markdown, no bullet characters
- All content must be technically accurate and professionally written
- Return ONLY the JSON object, nothing else
"""

    response = requests.post(
        f"{GEMINI_API_URL}?key={api_key}",
        headers={"Content-Type": "application/json"},
        json={
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.7,
                "maxOutputTokens": 4000,
            },
        },
        timeout=120,
    )
    response.raise_for_status()

    raw = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()

    # Strip any accidental markdown fences
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw).strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        start = raw.find("{")
        end   = raw.rfind("}") + 1
        if start != -1 and end > start:
            return json.loads(raw[start:end])
        raise RuntimeError("Could not parse Gemini response as JSON.") from None


# ═══════════════════════════════════════════
#  DOCX HELPERS
# ═══════════════════════════════════════════
def _xml_spacing(para, before=0, after=0, line=276, rule="auto"):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    sp.set(qn("w:line"),   str(line))
    sp.set(qn("w:lineRule"), rule)
    pPr.append(sp)


def _run(para, text, bold=False, italic=False, size=12,
         color=None, font=BODY_FONT, underline=False, superscript=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    r.font.name = font
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if superscript:
        r.font.superscript = True
    return r


def _body_para(doc, text, size=12, before=0, after=120, line=276):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _xml_spacing(p, before=before, after=after, line=line)
    _run(p, text, size=size)
    return p


def _heading(doc, text, size=13, before=200, after=100):
    """Red bold section heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _xml_spacing(p, before=before, after=after, line=240)
    _run(p, text, bold=True, size=size, color=RED, font=HEAD_FONT)
    return p


def _figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _xml_spacing(p, before=60, after=160, line=240)
    _run(p, text, bold=True, size=11, font=BODY_FONT)


def _placeholder_box(doc, caption):
    """Grey placeholder box for diagrams when no real image is provided."""
    try:
        from PIL import Image, ImageDraw
        W, H = 560, 280
        img  = Image.new("RGB", (W, H), (230, 230, 230))
        draw = ImageDraw.Draw(img)
        draw.rectangle([4, 4, W - 5, H - 5], outline=(160, 160, 160), width=2)
        lines = [caption[i:i+55] for i in range(0, min(len(caption), 110), 55)]
        y = H // 2 - len(lines) * 9
        for line in lines:
            draw.text((W // 2 - len(line) * 3, y), line, fill=(90, 90, 90))
            y += 18
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _xml_spacing(p, before=80, after=60, line=240)
        p.add_run().add_picture(buf, width=Inches(4.0))
    except Exception:
        p = doc.add_paragraph()
        _xml_spacing(p, before=80, after=60, line=240)


def _references_page(doc, references: list):
    _heading(doc, "References", before=160, after=80)
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Hanging indent
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"),    "720")
        ind.set(qn("w:hanging"), "720")
        pPr.append(ind)
        _xml_spacing(p, before=0, after=140, line=240)
        _run(p, f"{i}.\u00a0 {ref}", size=11)


# ═══════════════════════════════════════════
#  COVER PAGE  (with logo in top-right)
# ═══════════════════════════════════════════
def _build_cover(doc, title, author, university, college, department, grade, logo_bytes):
    """
    Mirror the exact layout of the original PDF cover:
      • Left column  : Ministry → University → College → Department → Grade  (bold, some red)
      • Right column : University logo image
      • Centre       : Big title
      • Bottom       : SUBMITTED BY  /  author name
    """

    # ── top two-column row: text LEFT  |  logo RIGHT ────────────────────────
    #    We use a single-row, two-cell table (invisible borders)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Table Grid"

    # Remove all borders
    for row in tbl.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "none")
                b.set(qn("w:sz"),    "0")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                tcBorders.append(b)
            tcPr.append(tcBorders)

    # Content width in DXA  (1 inch = 1440 DXA)
    # A4 21cm  minus  left 3cm + right 2.54cm  = 15.46 cm  ≈ 8755 DXA
    LEFT_W  = 6600   # ~4.6 inches
    RIGHT_W = 2155   # ~1.5 inches (logo cell)

    def _set_col_width(cell, dxa):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcW")):
            tcPr.remove(old)
        w = OxmlElement("w:tcW")
        w.set(qn("w:w"),    str(dxa))
        w.set(qn("w:type"), "dxa")
        tcPr.append(w)

    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    _set_col_width(left_cell,  LEFT_W)
    _set_col_width(right_cell, RIGHT_W)

    # Vertical align top
    for cell in (left_cell, right_cell):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        va   = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "top")
        tcPr.append(va)

    # ── LEFT cell — institution info ─────────────────────────────────────────
    def _lp(cell, text, bold=False, color=None, size=12, before=0, after=50):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _xml_spacing(p, before=before, after=after, line=240)
        _run(p, text, bold=bold, size=size, color=color, font=HEAD_FONT)
        return p

    # Remove the default empty paragraph that python-docx adds
    for old_p in left_cell.paragraphs:
        old_p._element.getparent().remove(old_p._element)

    _lp(left_cell, "Ministry of Higher Education",  bold=True, color=RED, size=13, after=0)
    _lp(left_cell, "and Scientific Research",        bold=True, color=RED, size=13, after=60)
    _lp(left_cell, university,   bold=True, size=12, after=40)
    _lp(left_cell, college,      bold=True, size=12, after=40)
    _lp(left_cell, department,   bold=True, size=12, after=40)

    # Grade with superscript
    g_num, g_sup, g_rest = _parse_grade(grade)
    p_grade = left_cell.add_paragraph()
    p_grade.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _xml_spacing(p_grade, before=0, after=0, line=240)
    _run(p_grade, g_num,  bold=True, size=12, font=HEAD_FONT)
    _run(p_grade, g_sup,  bold=True, size=10, font=HEAD_FONT, superscript=True)
    _run(p_grade, g_rest, bold=True, size=12, font=HEAD_FONT)

    # ── RIGHT cell — logo ────────────────────────────────────────────────────
    for old_p in right_cell.paragraphs:
        old_p._element.getparent().remove(old_p._element)

    p_logo = right_cell.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _xml_spacing(p_logo, before=0, after=0, line=240)

    if logo_bytes:
        try:
            p_logo.add_run().add_picture(BytesIO(logo_bytes), width=Inches(1.3))
        except Exception:
            _run(p_logo, "[logo]", size=9, color=RGBColor(0x80, 0x80, 0x80))
    else:
        _run(p_logo, "[logo]", size=9, color=RGBColor(0x80, 0x80, 0x80))

    # ── Spacer ───────────────────────────────────────────────────────────────
    sp = doc.add_paragraph()
    _xml_spacing(sp, before=0, after=0, line=240)
    sp.add_run(" ")

    # ── Big centred title ────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _xml_spacing(p_title, before=400, after=400, line=276)
    _run(p_title, title, bold=True, size=22, font=HEAD_FONT)

    # ── SUBMITTED BY ─────────────────────────────────────────────────────────
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _xml_spacing(p_sub, before=300, after=60, line=240)
    _run(p_sub, "SUBMITTED BY", bold=True, size=12,
         color=RED, font=HEAD_FONT, underline=True)

    # ── Author ───────────────────────────────────────────────────────────────
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _xml_spacing(p_auth, before=0, after=0, line=240)
    _run(p_auth, author, size=12, font=BODY_FONT)

    # ── Page break ───────────────────────────────────────────────────────────
    doc.add_page_break()


def _parse_grade(grade: str):
    """
    Split e.g. '3rd Grade' → ('3', 'rd', ' Grade')
    or '2nd Year' → ('2', 'nd', ' Year')
    Falls back to returning the whole string as-is.
    """
    import re
    m = re.match(r"^(\d+)(st|nd|rd|th)(.*)$", grade.strip(), re.IGNORECASE)
    if m:
        return m.group(1), m.group(2).lower(), m.group(3) or " Grade"
    return grade, "", ""


# ═══════════════════════════════════════════
#  MAIN  ── Build the full document
# ═══════════════════════════════════════════
def build_document(
    title: str,
    author: str,
    university: str,
    college: str,
    department: str,
    grade: str,
    content: dict,
    logo_bytes: bytes | None,
    output_path: str,
):
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.page_width    = Cm(PAGE_W)
        sec.page_height   = Cm(PAGE_H)
        sec.top_margin    = Cm(MARGIN_TOP)
        sec.bottom_margin = Cm(MARGIN_BOTTOM)
        sec.left_margin   = Cm(MARGIN_LEFT)
        sec.right_margin  = Cm(MARGIN_RIGHT)

    # ── Default style ──────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)

    # ── Cover page ─────────────────────────────────────────────────────────
    _build_cover(doc, title, author, university, college, department, grade, logo_bytes)

    # ── Introduction ───────────────────────────────────────────────────────
    _heading(doc, "Introduction", before=0, after=100)
    _body_para(doc, content["introduction"])

    # ── Sections ───────────────────────────────────────────────────────────
    for sec in content.get("sections", []):
        _heading(doc, sec["heading"])
        _body_para(doc, sec["body"])
        if sec.get("figure_caption"):
            _placeholder_box(doc, sec["figure_caption"])
            _figure_caption(doc, sec["figure_caption"])

    # ── Conclusion ─────────────────────────────────────────────────────────
    _heading(doc, "Conclusion")
    _body_para(doc, content["conclusion"])

    # ── References ─────────────────────────────────────────────────────────
    _references_page(doc, content.get("references", []))

    doc.save(output_path)
